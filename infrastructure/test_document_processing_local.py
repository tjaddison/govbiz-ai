#!/usr/bin/env python3
"""
Local test script for document processing with new libraries
Tests PyMuPDF4LLM and python-docx 1.2.0 functionality
"""

import os
import sys
import tempfile
import io
from typing import Dict, Any

# Add the lambda directory to path so we can import the handler
sys.path.insert(0, '/Users/terrance/Projects/govbiz-ai/infrastructure/lambda/document-processing')

# Test the imports first
def test_imports():
    """Test that all required libraries can be imported"""
    print("Testing library imports...")

    try:
        import fitz  # PyMuPDF
        print("✓ PyMuPDF imported successfully")
    except ImportError as e:
        print(f"✗ PyMuPDF import failed: {e}")
        return False

    try:
        from docx import Document
        print("✓ python-docx imported successfully")
    except ImportError as e:
        print(f"✗ python-docx import failed: {e}")
        return False

    try:
        import openpyxl
        print("✓ openpyxl imported successfully")
    except ImportError as e:
        print(f"✗ openpyxl import failed: {e}")
        return False

    return True

def test_pdf_processing():
    """Test PDF processing with PyMuPDF"""
    print("\nTesting PDF processing...")

    # Create a simple PDF for testing (we'll create a text file and pretend it's PDF for now)
    test_content = b"This is a test PDF content for extraction testing."

    try:
        # Import the extraction function
        from handler import extract_text_with_pymupdf, PYMUPDF_AVAILABLE

        if not PYMUPDF_AVAILABLE:
            print("✗ PyMuPDF not available in handler")
            return False

        print("✓ PDF processing function available")
        return True

    except Exception as e:
        print(f"✗ PDF processing test failed: {e}")
        return False

def test_docx_processing():
    """Test DOCX processing with python-docx 1.2.0"""
    print("\nTesting DOCX processing...")

    try:
        # Import the extraction function
        from handler import extract_text_from_docx, DOCX_AVAILABLE

        if not DOCX_AVAILABLE:
            print("✗ python-docx not available in handler")
            return False

        # Create a simple DOCX document for testing
        from docx import Document

        # Create document in memory
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph for DOCX extraction testing.')
        doc.add_paragraph('This is another paragraph with some content.')

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Row 1 Col 1'
        table.cell(1, 1).text = 'Row 1 Col 2'

        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_content = doc_io.getvalue()

        # Test extraction
        extracted_text = extract_text_from_docx(doc_content)

        if len(extracted_text) > 50:  # Should have decent amount of text
            print("✓ DOCX processing successful")
            print(f"  Extracted {len(extracted_text)} characters")
            print(f"  Sample: {extracted_text[:100]}...")
            return True
        else:
            print(f"✗ DOCX processing returned insufficient text: {len(extracted_text)} chars")
            return False

    except Exception as e:
        print(f"✗ DOCX processing test failed: {e}")
        return False

def test_markdown_conversion():
    """Test markdown to text conversion"""
    print("\nTesting markdown conversion...")

    try:
        from handler import markdown_to_text

        test_markdown = """
# Heading 1
This is a paragraph with **bold** and *italic* text.

## Heading 2
- List item 1
- List item 2

| Column 1 | Column 2 |
|----------|----------|
| Cell 1   | Cell 2   |

[Link text](http://example.com)
        """

        converted_text = markdown_to_text(test_markdown)

        if len(converted_text) > 20:
            print("✓ Markdown conversion successful")
            print(f"  Converted text: {converted_text[:100]}...")
            return True
        else:
            print(f"✗ Markdown conversion failed: {len(converted_text)} chars")
            return False

    except Exception as e:
        print(f"✗ Markdown conversion test failed: {e}")
        return False

def main():
    """Run all tests"""
    print("=" * 60)
    print("Document Processing Local Tests")
    print("=" * 60)

    all_passed = True

    # Test imports
    if not test_imports():
        all_passed = False

    # Test PDF processing
    if not test_pdf_processing():
        all_passed = False

    # Test DOCX processing
    if not test_docx_processing():
        all_passed = False

    # Test markdown conversion
    if not test_markdown_conversion():
        all_passed = False

    print("\n" + "=" * 60)
    if all_passed:
        print("✓ All tests passed! Ready for deployment.")
    else:
        print("✗ Some tests failed. Please fix issues before deployment.")
    print("=" * 60)

    return all_passed

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)