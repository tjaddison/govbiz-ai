import json
import boto3
import os
from typing import Dict, Any, List
import logging
from datetime import datetime
import uuid
import io
from urllib.parse import unquote
import re
import tempfile

# Import new libraries for document processing
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logger = logging.getLogger()
logger.setLevel(logging.INFO)

s3_client = boto3.client('s3')
textract_client = boto3.client('textract')
bedrock_client = boto3.client('bedrock-runtime')
dynamodb = boto3.resource('dynamodb')
sqs = boto3.client('sqs')

DOCUMENTS_BUCKET = os.environ['DOCUMENTS_BUCKET']
EMBEDDINGS_BUCKET = os.environ['EMBEDDINGS_BUCKET']
COMPANIES_TABLE = os.environ['COMPANIES_TABLE']
PROCESSING_QUEUE_URL = os.environ.get('PROCESSING_QUEUE_URL')

def lambda_handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Process uploaded documents: extract text, categorize, generate embeddings
    Triggered by S3 upload events or SQS messages
    """
    try:
        # Handle S3 trigger or SQS message
        if 'Records' in event:
            for record in event['Records']:
                if record.get('eventSource') == 'aws:s3':
                    # S3 trigger
                    bucket = record['s3']['bucket']['name']
                    key = unquote(record['s3']['object']['key'])
                    process_document_from_s3(bucket, key)
                elif record.get('eventSource') == 'aws:sqs':
                    # SQS message
                    body = json.loads(record['body'])
                    process_document_from_message(body)
        else:
            # Direct Lambda invoke
            bucket = event.get('bucket')
            key = event.get('key')
            if bucket and key:
                process_document_from_s3(bucket, key)
            else:
                logger.error("Invalid event format")
                return {'statusCode': 400, 'body': 'Invalid event format'}

        return {'statusCode': 200, 'body': 'Processing completed'}

    except Exception as e:
        logger.error(f"Document processing error: {str(e)}")
        return {'statusCode': 500, 'body': f'Processing failed: {str(e)}'}

def process_document_from_s3(bucket: str, key: str):
    """Process document from S3 upload event"""
    try:
        # Extract company_id and document_id from S3 key
        # Expected format: {company_id}/raw/{document_id}/{filename}
        parts = key.split('/')
        if len(parts) < 4:
            logger.error(f"Invalid S3 key format: {key}")
            return

        company_id = parts[0]
        document_id = parts[2]
        filename = parts[3]

        process_document(company_id, document_id, bucket, key, filename)

    except Exception as e:
        logger.error(f"Error processing document from S3: {str(e)}")

def process_document_from_message(message: Dict[str, Any]):
    """Process document from SQS message"""
    try:
        company_id = message.get('company_id')
        document_id = message.get('document_id')
        bucket = message.get('bucket')
        key = message.get('key')
        filename = message.get('filename')

        if not all([company_id, document_id, bucket, key, filename]):
            logger.error(f"Missing required fields in message: {message}")
            return

        process_document(company_id, document_id, bucket, key, filename)

    except Exception as e:
        logger.error(f"Error processing document from message: {str(e)}")

def process_document(company_id: str, document_id: str, bucket: str, key: str, filename: str):
    """Main document processing pipeline"""
    try:
        logger.info(f"Processing document: {filename} for company: {company_id}")

        # Download document from S3
        response = s3_client.get_object(Bucket=bucket, Key=key)
        document_content = response['Body'].read()

        # Extract text based on file type
        file_extension = os.path.splitext(filename)[1].lower()
        extracted_text = extract_text_from_document(document_content, file_extension, filename)

        if not extracted_text or len(extracted_text.strip()) < 50:
            logger.warning(f"Insufficient text extracted from {filename}")
            update_document_status(company_id, document_id, 'failed', 'Insufficient text content')
            return

        # Categorize document
        category = categorize_document(extracted_text, filename)

        # Generate embeddings
        embeddings_metadata = generate_document_embeddings(extracted_text, company_id, document_id, category)

        # Store processed text and metadata
        processed_key = f"{company_id}/processed/{document_id}/{filename}.txt"
        s3_client.put_object(
            Bucket=DOCUMENTS_BUCKET,
            Key=processed_key,
            Body=extracted_text.encode('utf-8'),
            ContentType='text/plain'
        )

        # Update document status in DynamoDB
        update_document_status(
            company_id,
            document_id,
            'processed',
            'Successfully processed',
            {
                'processed_key': processed_key,
                'category': category,
                'text_length': len(extracted_text),
                'embeddings_count': len(embeddings_metadata),
                'embeddings_ids': [em['embedding_id'] for em in embeddings_metadata]
            }
        )

        # Trigger company profile re-embedding
        trigger_profile_reembedding(company_id)

        logger.info(f"Successfully processed document {document_id} for company {company_id}")

    except Exception as e:
        logger.error(f"Error processing document {document_id}: {str(e)}")
        update_document_status(company_id, document_id, 'failed', str(e))

def extract_text_from_document(content: bytes, file_extension: str, filename: str) -> str:
    """Extract text from document using appropriate method based on file type"""
    try:
        if file_extension == '.pdf':
            try:
                return extract_text_with_pymupdf(content)
            except Exception as pdf_error:
                logger.warning(f"PyMuPDF failed for {filename}: {str(pdf_error)}")
                # Fallback to Textract for PDFs
                logger.info(f"Falling back to Textract for PDF: {filename}")
                return extract_text_with_textract(content, filename)

        elif file_extension == '.docx':
            try:
                return extract_text_from_docx(content)
            except Exception as docx_error:
                logger.warning(f"python-docx failed for {filename}: {str(docx_error)}")
                # For DOCX files, try Textract as fallback (though limited support)
                logger.info(f"Falling back to Textract for DOCX: {filename}")
                return extract_text_with_textract(content, filename)

        elif file_extension == '.doc':
            # Legacy .doc files - try Textract first, then give helpful message
            logger.warning(f"Legacy .doc format detected: {filename}")
            try:
                return extract_text_with_textract(content, filename)
            except Exception as doc_error:
                logger.error(f"Textract failed for .doc file {filename}: {str(doc_error)}")
                return f"Legacy .doc format extraction failed. Please convert to .docx format for better results. File: {filename}"

        elif file_extension in ['.xlsx', '.xls']:
            return extract_text_from_excel(content)

        elif file_extension in ['.txt']:
            return content.decode('utf-8', errors='ignore')

        elif file_extension.lower() in ['.png', '.jpg', '.jpeg', '.tiff', '.tif']:
            # Image files - use Textract OCR
            logger.info(f"Processing image file with Textract OCR: {filename}")
            return extract_text_with_textract(content, filename)

        else:
            # Unknown format - try Textract as last resort
            logger.warning(f"Unknown file format {file_extension}, trying Textract: {filename}")
            return extract_text_with_textract(content, filename)

    except Exception as e:
        logger.error(f"All text extraction methods failed for {filename}: {str(e)}")
        return ""

def extract_text_with_pymupdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF with improved extraction"""
    if not PYMUPDF_AVAILABLE:
        logger.error("PyMuPDF library not available")
        raise Exception("PyMuPDF library not installed")

    try:
        # Create a temporary file to work with PyMuPDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(content)
            temp_file.flush()

            try:
                # Open the PDF document
                doc = fitz.open(temp_file.name)
                text_parts = []

                # Extract text from each page with improved formatting
                for page_num in range(doc.page_count):
                    page = doc[page_num]

                    # Get page text with layout preservation
                    page_text = page.get_text("text")

                    if page_text.strip():
                        # Add page separator for multi-page documents
                        if page_num > 0:
                            text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                        text_parts.append(page_text.strip())

                doc.close()

                full_text = '\n'.join(text_parts)
                logger.info(f"PyMuPDF extracted {len(full_text)} characters from {doc.page_count} pages")

                return clean_extracted_text(full_text)

            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file.name)
                except:
                    pass  # Ignore cleanup errors

    except Exception as e:
        logger.error(f"PyMuPDF extraction failed: {str(e)}")
        raise

def extract_text_with_textract(content: bytes, filename: str) -> str:
    """Extract text using Amazon Textract (async for large documents)"""
    try:
        # For documents under 5MB, use synchronous detection
        if len(content) < 5 * 1024 * 1024:
            response = textract_client.detect_document_text(
                Document={'Bytes': content}
            )

            text_parts = []
            for block in response.get('Blocks', []):
                if block['BlockType'] == 'LINE':
                    text_parts.append(block['Text'])

            return clean_extracted_text('\n'.join(text_parts))
        else:
            # For larger documents, use asynchronous processing
            return extract_text_with_textract_async(content, filename)

    except Exception as e:
        logger.error(f"Textract extraction failed: {str(e)}")
        raise

def extract_text_with_textract_async(content: bytes, filename: str) -> str:
    """Handle large document processing with Textract async API"""
    try:
        # Upload to temporary S3 location for Textract
        temp_key = f"temp/textract/{uuid.uuid4()}/{filename}"
        s3_client.put_object(
            Bucket=DOCUMENTS_BUCKET,
            Key=temp_key,
            Body=content
        )

        # Start async job
        response = textract_client.start_document_text_detection(
            DocumentLocation={
                'S3Object': {
                    'Bucket': DOCUMENTS_BUCKET,
                    'Name': temp_key
                }
            }
        )

        job_id = response['JobId']

        # Poll for completion (simplified - in production, use SQS notifications)
        import time
        max_attempts = 30
        attempt = 0

        while attempt < max_attempts:
            result = textract_client.get_document_text_detection(JobId=job_id)
            status = result['JobStatus']

            if status == 'SUCCEEDED':
                text_parts = []
                for block in result.get('Blocks', []):
                    if block['BlockType'] == 'LINE':
                        text_parts.append(block['Text'])

                # Clean up temp file
                s3_client.delete_object(Bucket=DOCUMENTS_BUCKET, Key=temp_key)

                return clean_extracted_text('\n'.join(text_parts))
            elif status == 'FAILED':
                logger.error(f"Textract job failed: {job_id}")
                break
            else:
                time.sleep(5)
                attempt += 1

        # Clean up temp file
        s3_client.delete_object(Bucket=DOCUMENTS_BUCKET, Key=temp_key)
        raise Exception(f"Textract job did not complete successfully: {job_id}")

    except Exception as e:
        logger.error(f"Textract async extraction failed: {str(e)}")
        raise

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from Word (.docx) document using python-docx 1.2.0"""
    if not DOCX_AVAILABLE:
        logger.error("python-docx library not available")
        raise Exception("python-docx library not installed")

    try:
        # Create a Document object from bytes
        doc_io = io.BytesIO(content)
        document = Document(doc_io)

        text_parts = []

        # Extract text from all paragraphs with better formatting preservation
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                # Check for heading styles to preserve document structure
                if paragraph.style.name.startswith('Heading'):
                    text_parts.append(f"\n{paragraph.text.strip()}\n")
                else:
                    text_parts.append(paragraph.text.strip())

        # Extract text from tables with better structure preservation
        for table in document.tables:
            table_text = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    # Extract text from each cell, including nested paragraphs
                    cell_text = []
                    for cell_paragraph in cell.paragraphs:
                        if cell_paragraph.text.strip():
                            cell_text.append(cell_paragraph.text.strip())

                    if cell_text:
                        row_cells.append(' '.join(cell_text))

                if row_cells:
                    table_text.append(' | '.join(row_cells))

            if table_text:
                text_parts.append('\n' + '\n'.join(table_text) + '\n')

        # Extract text from headers and footers if present
        try:
            for section in document.sections:
                # Header text
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[Header: {paragraph.text.strip()}]")

                # Footer text
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[Footer: {paragraph.text.strip()}]")
        except Exception as header_footer_error:
            # Headers/footers extraction can fail, but don't let it break the whole process
            logger.warning(f"Could not extract headers/footers: {str(header_footer_error)}")

        # Join all text parts
        full_text = '\n'.join(text_parts)

        # Clean up excessive line breaks but preserve paragraph structure
        full_text = re.sub(r'\n{3,}', '\n\n', full_text)

        logger.info(f"Extracted {len(full_text)} characters from Word document")
        return clean_extracted_text(full_text)

    except Exception as e:
        logger.error(f"Error extracting text from Word document: {str(e)}")
        raise

def extract_text_from_excel(content: bytes) -> str:
    """Extract text from Excel (.xlsx/.xls) document"""
    if not EXCEL_AVAILABLE:
        logger.error("openpyxl library not available")
        raise Exception("openpyxl library not installed")

    try:
        # Load workbook from bytes
        excel_io = io.BytesIO(content)
        workbook = load_workbook(excel_io, read_only=True, data_only=True)

        text_parts = []

        # Extract text from all worksheets
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]

            # Add sheet name as header
            text_parts.append(f"=== {sheet_name} ===")

            # Extract cell values
            for row in worksheet.iter_rows(values_only=True):
                row_text = []
                for cell_value in row:
                    if cell_value is not None:
                        row_text.append(str(cell_value).strip())

                if row_text:  # Only add non-empty rows
                    text_parts.append(' | '.join(row_text))

        full_text = '\n'.join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from Excel document")
        return clean_extracted_text(full_text)

    except Exception as e:
        logger.error(f"Error extracting text from Excel document: {str(e)}")
        raise

def markdown_to_text(markdown: str) -> str:
    """Convert markdown to clean text while preserving structure"""
    if not markdown:
        return ""

    # Remove markdown formatting while keeping structure
    text = markdown

    # Remove markdown headers but keep the text
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)

    # Remove markdown bold/italic
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*([^*]+)\*', r'\1', text)      # Italic
    text = re.sub(r'__([^_]+)__', r'\1', text)      # Bold alt
    text = re.sub(r'_([^_]+)_', r'\1', text)        # Italic alt

    # Remove markdown links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    # Remove markdown code blocks
    text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # Convert markdown lists to simple format
    text = re.sub(r'^\s*[-*+]\s+', '• ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '• ', text, flags=re.MULTILINE)

    # Remove markdown tables formatting but keep content
    text = re.sub(r'\|', ' ', text)
    text = re.sub(r'^[-:\s|]+$', '', text, flags=re.MULTILINE)

    return text.strip()

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""

    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)

    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\=\|]', '', text)

    # Remove common PDF artifacts
    text = re.sub(r'Page \d+ of \d+', '', text)
    text = re.sub(r'Confidential.*?(?=\n|\r\n|\r)', '', text, flags=re.IGNORECASE)

    # Normalize line breaks
    text = re.sub(r'\r\n|\r', '\n', text)

    return text.strip()

def categorize_document(text: str, filename: str) -> str:
    """Automatically categorize document based on content and filename"""
    try:
        text_lower = text.lower()
        filename_lower = filename.lower()

        # Score-based categorization for better accuracy
        category_scores = {
            'capability_statement': 0,
            'resume': 0,
            'past_performance': 0,
            'certification': 0,
            'financial': 0,
            'proposal': 0,
            'other': 0
        }

        # Capability Statement indicators
        capability_terms = ['capability statement', 'capabilities', 'core competencies', 'company overview', 'business profile', 'corporate overview']
        capability_filename_terms = ['capability', 'cap', 'overview', 'profile']
        for term in capability_terms:
            if term in text_lower:
                category_scores['capability_statement'] += 3
        for term in capability_filename_terms:
            if term in filename_lower:
                category_scores['capability_statement'] += 2

        # Resume indicators (stronger filename weighting)
        resume_text_terms = ['resume', 'curriculum vitae', 'cv', 'work experience', 'professional experience', 'education', 'skills', 'employment history']
        resume_filename_terms = ['resume', 'cv', 'bio']
        for term in resume_text_terms:
            if term in text_lower:
                category_scores['resume'] += 2
        for term in resume_filename_terms:
            if term in filename_lower:
                category_scores['resume'] += 4  # Strong filename indicator

        # Past Performance indicators
        past_perf_terms = ['past performance', 'contract', 'project', 'cpars', 'award', 'client work', 'project history', 'contract performance']
        past_perf_filename_terms = ['performance', 'contract', 'project', 'award']
        for term in past_perf_terms:
            if term in text_lower:
                category_scores['past_performance'] += 3
        for term in past_perf_filename_terms:
            if term in filename_lower:
                category_scores['past_performance'] += 2

        # Certification indicators
        cert_terms = ['certification', 'certificate', '8(a)', 'wosb', 'sdvosb', 'hubzone', 'certified', 'certification authority', 'iso']
        cert_filename_terms = ['cert', 'certification', '8a', 'wosb', 'sdvosb']
        for term in cert_terms:
            if term in text_lower:
                category_scores['certification'] += 3
        for term in cert_filename_terms:
            if term in filename_lower:
                category_scores['certification'] += 2

        # Financial document indicators
        financial_terms = ['financial', 'income', 'revenue', 'balance sheet', 'cash flow', 'profit loss', 'financial statement', 'accounting']
        financial_filename_terms = ['financial', 'finance', 'statement', 'accounting']
        for term in financial_terms:
            if term in text_lower:
                category_scores['financial'] += 3
        for term in financial_filename_terms:
            if term in filename_lower:
                category_scores['financial'] += 2

        # Proposal indicators
        proposal_terms = ['proposal', 'response', 'rfp', 'solicitation', 'bid response', 'technical proposal', 'cost proposal']
        proposal_filename_terms = ['proposal', 'response', 'rfp', 'bid']
        for term in proposal_terms:
            if term in text_lower:
                category_scores['proposal'] += 3
        for term in proposal_filename_terms:
            if term in filename_lower:
                category_scores['proposal'] += 2

        # Find the category with the highest score
        max_score = max(category_scores.values())
        if max_score >= 2:  # Minimum confidence threshold
            best_category = max(category_scores, key=category_scores.get)
            logger.info(f"Categorized document as '{best_category}' with score {max_score}")
            return best_category
        else:
            logger.info(f"Document categorized as 'other' - no strong indicators found")
            return 'other'

    except Exception as e:
        logger.warning(f"Document categorization failed: {str(e)}")
        return 'other'

def generate_document_embeddings(text: str, company_id: str, document_id: str, category: str) -> List[Dict]:
    """Generate embeddings for document using Amazon Bedrock"""
    try:
        # Chunk the text into manageable pieces
        chunks = chunk_text(text, max_tokens=1000, overlap_tokens=200)

        embeddings_metadata = []

        for i, chunk in enumerate(chunks):
            # Generate embedding using Bedrock Titan
            embedding_response = bedrock_client.invoke_model(
                modelId='amazon.titan-embed-text-v2:0',
                body=json.dumps({
                    'inputText': chunk,
                    'dimensions': 1024
                })
            )

            embedding_result = json.loads(embedding_response['body'].read())
            embedding_vector = embedding_result['embedding']

            # Store embedding in S3 with metadata
            embedding_id = f"{document_id}_chunk_{i}"
            embedding_key = f"{company_id}/embeddings/documents/{embedding_id}.json"

            embedding_data = {
                'embedding_id': embedding_id,
                'company_id': company_id,
                'document_id': document_id,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'category': category,
                'text_chunk': chunk,
                'embedding': embedding_vector,
                'created_at': datetime.utcnow().isoformat() + 'Z'
            }

            s3_client.put_object(
                Bucket=EMBEDDINGS_BUCKET,
                Key=embedding_key,
                Body=json.dumps(embedding_data),
                ContentType='application/json'
            )

            embeddings_metadata.append({
                'embedding_id': embedding_id,
                'embedding_key': embedding_key,
                'chunk_index': i,
                'chunk_text_length': len(chunk)
            })

            logger.info(f"Generated embedding {embedding_id} for document {document_id}")

        return embeddings_metadata

    except Exception as e:
        logger.error(f"Error generating embeddings for document {document_id}: {str(e)}")
        raise

def chunk_text(text: str, max_tokens: int = 1000, overlap_tokens: int = 200) -> List[str]:
    """Split text into overlapping chunks for embedding generation"""
    if not text:
        return []

    # Simple word-based chunking (in production, use proper tokenization)
    words = text.split()
    chunks = []

    # Estimate tokens as roughly 0.75 * words
    words_per_chunk = int(max_tokens * 0.75)
    overlap_words = int(overlap_tokens * 0.75)

    start = 0
    while start < len(words):
        end = min(start + words_per_chunk, len(words))
        chunk_words = words[start:end]
        chunks.append(' '.join(chunk_words))

        if end >= len(words):
            break

        start = end - overlap_words

    return chunks

def update_document_status(company_id: str, document_id: str, status: str, message: str, metadata: Dict = None):
    """Update document processing status in DynamoDB"""
    try:
        companies_table = dynamodb.Table(COMPANIES_TABLE)

        # Get current company profile
        response = companies_table.get_item(Key={'company_id': company_id})

        if 'Item' not in response:
            logger.error(f"Company not found: {company_id}")
            return

        documents = response['Item'].get('documents', [])

        # Find and update the specific document
        for i, doc in enumerate(documents):
            if doc.get('document_id') == document_id:
                documents[i]['status'] = status
                documents[i]['status_message'] = message
                documents[i]['updated_at'] = datetime.utcnow().isoformat() + 'Z'

                if metadata:
                    documents[i].update(metadata)

                break

        # Update the company profile
        companies_table.update_item(
            Key={'company_id': company_id},
            UpdateExpression="SET documents = :docs, updated_at = :updated_at",
            ExpressionAttributeValues={
                ':docs': documents,
                ':updated_at': datetime.utcnow().isoformat() + 'Z'
            }
        )

    except Exception as e:
        logger.error(f"Error updating document status: {str(e)}")

def trigger_profile_reembedding(company_id: str):
    """Trigger company profile re-embedding after document processing"""
    try:
        profile_embedding_queue_url = os.environ.get('PROFILE_EMBEDDING_QUEUE_URL')
        if profile_embedding_queue_url:
            sqs.send_message(
                QueueUrl=profile_embedding_queue_url,
                MessageBody=json.dumps({
                    'action': 'reembed_profile',
                    'company_id': company_id,
                    'timestamp': datetime.utcnow().isoformat() + 'Z'
                })
            )
            logger.info(f"Triggered profile re-embedding for company: {company_id}")
        else:
            logger.warning("PROFILE_EMBEDDING_QUEUE_URL not configured")
    except Exception as e:
        logger.warning(f"Failed to trigger profile re-embedding: {str(e)}")